import os
import base64
from PIL import Image
import PyPDF2
from docx import Document
import io

class FileProcessor:
    """文件處理器，支援多種文件格式"""
    
    @staticmethod
    def process_image(image_file):
        """處理圖片文件"""
        try:
            # 讀取圖片
            image = Image.open(image_file)
            
            # 轉換為base64
            buffered = io.BytesIO()
            image.save(buffered, format=image.format or 'PNG')
            img_str = base64.b64encode(buffered.getvalue()).decode()
            
            # 獲取圖片信息
            width, height = image.size
            format_info = image.format or 'Unknown'
            
            return {
                'type': 'image',
                'content': img_str,
                'info': f"圖片格式: {format_info}, 尺寸: {width}x{height}",
                'filename': image_file.name
            }
        except Exception as e:
            return {'error': f"圖片處理錯誤: {str(e)}"}
    
    @staticmethod
    def process_pdf(pdf_file):
        """處理PDF文件"""
        try:
            # 讀取PDF
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # 提取文字內容
            text_content = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text_content += f"\n--- 第 {page_num + 1} 頁 ---\n"
                text_content += page.extract_text()
            
            return {
                'type': 'pdf',
                'content': text_content,
                'info': f"PDF頁數: {len(pdf_reader.pages)}",
                'filename': pdf_file.name
            }
        except Exception as e:
            return {'error': f"PDF處理錯誤: {str(e)}"}
    
    @staticmethod
    def process_docx(docx_file):
        """處理Word文檔"""
        try:
            # 讀取Word文檔
            doc = Document(docx_file)
            
            # 提取文字內容
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            return {
                'type': 'docx',
                'content': text_content,
                'info': f"Word文檔段落數: {len(doc.paragraphs)}",
                'filename': docx_file.name
            }
        except Exception as e:
            return {'error': f"Word文檔處理錯誤: {str(e)}"}
    
    @staticmethod
    def process_text_file(text_file):
        """處理文字文件"""
        try:
            content = text_file.read().decode('utf-8')
            return {
                'type': 'text',
                'content': content,
                'info': f"文字文件大小: {len(content)} 字符",
                'filename': text_file.name
            }
        except Exception as e:
            return {'error': f"文字文件處理錯誤: {str(e)}"}
    
    @staticmethod
    def process_file(uploaded_file):
        """根據文件類型處理文件"""
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        
        # 圖片格式
        if file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']:
            return FileProcessor.process_image(uploaded_file)
        
        # PDF格式
        elif file_extension == '.pdf':
            return FileProcessor.process_pdf(uploaded_file)
        
        # Word文檔
        elif file_extension in ['.docx', '.doc']:
            return FileProcessor.process_docx(uploaded_file)
        
        # 文字文件
        elif file_extension in ['.txt', '.md', '.csv']:
            return FileProcessor.process_text_file(uploaded_file)
        
        else:
            return {'error': f"不支援的文件格式: {file_extension}"} 